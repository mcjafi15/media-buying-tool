# tests/test_exporters.py
from io import BytesIO
from docx import Document

DEAL = {
    "name": "Test Liquidation",
    "type": "industrial",
    "location": "Detroit, MI",
    "target_audience": "Manufacturing buyers",
    "total_budget": 20000.0,
    "start_date": "2026-07-01",
    "end_date": "2026-08-01",
}

PLAN = {
    "channel_mix": {
        "digital": {
            "pct": 60, "usd": 12000.0,
            "breakdown": {"google_search": 40, "google_display": 20, "meta": 25, "linkedin": 15},
        },
        "print": {"pct": 25, "usd": 5000.0, "notes": "Industry Week"},
        "ooh": {"pct": 15, "usd": 3000.0, "notes": "Near facility"},
    },
    "rationale": "Strong B2B digital reach.",
    "flight_schedule": "Heavy first 2 weeks.",
    "targeting_notes": "Plant managers, ops directors.",
    "keywords": ["equipment auction", "liquidation"],
}


def _load_docx(raw_bytes):
    return Document(BytesIO(raw_bytes))


def _full_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_generate_digital_brief_returns_bytes():
    from core.exporters import generate_digital_brief
    result = generate_digital_brief(DEAL, PLAN)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_digital_brief_contains_deal_name():
    from core.exporters import generate_digital_brief
    doc = _load_docx(generate_digital_brief(DEAL, PLAN))
    assert "Test Liquidation" in _full_text(doc)


def test_generate_print_brief_returns_bytes():
    from core.exporters import generate_print_brief
    result = generate_print_brief(DEAL, PLAN)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_ooh_brief_returns_bytes():
    from core.exporters import generate_ooh_brief
    result = generate_ooh_brief(DEAL, PLAN)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_rfp_template_returns_bytes():
    from core.exporters import generate_rfp_template
    result = generate_rfp_template(DEAL, PLAN, vendor_type="print")
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_budget_approval_form_returns_bytes():
    from core.exporters import generate_budget_approval_form
    result = generate_budget_approval_form(DEAL, PLAN)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_media_plan_summary_returns_bytes():
    from core.exporters import generate_media_plan_summary
    result = generate_media_plan_summary(DEAL, PLAN)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_media_plan_summary_contains_budget():
    from core.exporters import generate_media_plan_summary
    doc = _load_docx(generate_media_plan_summary(DEAL, PLAN))
    text = _full_text(doc)
    assert "20,000" in text or "20000" in text


def test_generate_performance_report_returns_pdf_bytes():
    from core.exporters import generate_performance_report
    snapshots = [
        {"platform": "google", "impressions": 10000, "clicks": 500, "spend": 4500.0,
         "conversions": 12, "synced_at": "2026-07-15"},
        {"platform": "meta", "impressions": 8000, "clicks": 300, "spend": 3000.0,
         "conversions": 8, "synced_at": "2026-07-15"},
    ]
    placements = [
        {"channel": "print", "vendor_name": "Industry Week", "contracted_cost": 5000.0,
         "actual_spend": 5000.0, "status": "complete"},
    ]
    result = generate_performance_report(DEAL, snapshots, placements)
    assert isinstance(result, bytes)
    assert result[:4] == b"%PDF"


def test_generate_performance_report_empty_data_returns_valid_pdf():
    from core.exporters import generate_performance_report
    result = generate_performance_report(DEAL, [], [])
    assert isinstance(result, bytes)
    assert result[:4] == b"%PDF"


CREATIVE = {
    "google_search": {
        "headlines": [f"Headline {i}" for i in range(1, 16)],
        "descriptions": ["Desc 1", "Desc 2", "Desc 3", "Desc 4"],
        "notes": "Target plant managers.",
    },
    "meta": {
        "primary_text": ["Text 1", "Text 2", "Text 3"],
        "headlines": ["Meta H1", "Meta H2", "Meta H3"],
        "cta": "Learn More",
        "notes": "Broad audience.",
    },
    "linkedin": {
        "headline": ["LI H1", "LI H2", "LI H3"],
        "body": ["LI Body 1", "LI Body 2"],
        "notes": "Target ops directors.",
    },
    "print": {
        "headlines": ["Print H1", "Print H2", "Print H3"],
        "body_short": "Short body copy here.",
        "body_long": "Long body copy here with more detail.",
        "cta": "Visit us online.",
        "notes": "Industry Week placement.",
    },
    "ooh": {
        "headlines": ["Big Sale Now", "Auction This Week", "Bid Today Win"],
        "subheads": ["Chicago July 2026", "Equipment Available"],
        "cta": "HilcoGlobal.com",
        "notes": "Near facility.",
    },
}


def test_generate_creative_brief_doc_returns_bytes():
    from core.exporters import generate_creative_brief_doc
    result = generate_creative_brief_doc(DEAL, PLAN, CREATIVE)
    assert isinstance(result, bytes) and len(result) > 0


def test_generate_creative_brief_doc_contains_deal_name_and_channels():
    from core.exporters import generate_creative_brief_doc
    doc = _load_docx(generate_creative_brief_doc(DEAL, PLAN, CREATIVE))
    text = _full_text(doc)
    assert "Test Liquidation" in text
    assert "Google" in text
    assert "Meta" in text
    assert "LinkedIn" in text or "Sponsored" in text
    assert "Print" in text
    assert "Out-of-Home" in text or "Billboard" in text or "OOH" in text or "Signage" in text
