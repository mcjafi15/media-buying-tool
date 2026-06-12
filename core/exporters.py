# core/exporters.py
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _new_doc(title: str) -> Document:
    doc = Document()
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return doc


def _bold_para(doc, label: str, value: str):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run(str(value) if value is not None else "")


def _deal_overview(doc, deal: dict):
    doc.add_heading("Deal Overview", 1)
    _bold_para(doc, "Deal Name", deal["name"])
    _bold_para(doc, "Deal Type", deal["type"].replace("_", " ").title())
    _bold_para(doc, "Location", deal["location"])
    _bold_para(doc, "Campaign Dates", f"{deal['start_date']} to {deal['end_date']}")
    _bold_para(doc, "Target Audience", deal.get("target_audience") or "See notes")


def _to_bytes(doc: Document) -> bytes:
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_digital_brief(deal: dict, plan: dict) -> bytes:
    doc = _new_doc(f"Digital Media Brief — {deal['name']}")
    _deal_overview(doc, deal)

    digital = plan["channel_mix"]["digital"]
    doc.add_heading("Digital Budget", 1)
    _bold_para(doc, "Total Digital Budget", f"${digital['usd']:,.0f} ({digital['pct']}% of total)")

    doc.add_heading("Platform Breakdown", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Platform", "% of Digital", "Est. Budget"
    for platform, pct in digital["breakdown"].items():
        sub = digital["usd"] * pct / 100
        row = table.add_row().cells
        row[0].text = platform.replace("_", " ").title()
        row[1].text = f"{pct}%"
        row[2].text = f"${sub:,.0f}"

    doc.add_heading("Targeting", 1)
    doc.add_paragraph(plan.get("targeting_notes", ""))
    if plan.get("keywords"):
        _bold_para(doc, "Keywords", ", ".join(plan["keywords"]))

    doc.add_heading("Flight Schedule", 1)
    doc.add_paragraph(plan.get("flight_schedule", ""))
    return _to_bytes(doc)


def generate_print_brief(deal: dict, plan: dict) -> bytes:
    doc = _new_doc(f"Print Media Brief — {deal['name']}")
    _deal_overview(doc, deal)

    print_data = plan["channel_mix"].get("print", {"pct": 0, "usd": 0.0, "notes": ""})
    doc.add_heading("Print Budget & Placements", 1)
    _bold_para(doc, "Total Print Budget", f"${print_data['usd']:,.0f} ({print_data['pct']}% of total)")
    _bold_para(doc, "Recommended Publications", print_data.get("notes", ""))
    doc.add_heading("Flight Schedule", 1)
    doc.add_paragraph(plan.get("flight_schedule", ""))
    doc.add_heading("Ad Specifications", 1)
    doc.add_paragraph("Please provide specifications for: full page, half page, and quarter page formats. "
                      "Include bleed dimensions, file format requirements, and submission deadlines.")
    return _to_bytes(doc)


def generate_ooh_brief(deal: dict, plan: dict) -> bytes:
    doc = _new_doc(f"Out-of-Home Media Brief — {deal['name']}")
    _deal_overview(doc, deal)

    ooh = plan["channel_mix"].get("ooh", {"pct": 0, "usd": 0.0, "notes": ""})
    doc.add_heading("OOH Budget & Placements", 1)
    _bold_para(doc, "Total OOH Budget", f"${ooh['usd']:,.0f} ({ooh['pct']}% of total)")
    _bold_para(doc, "Placement Guidance", ooh.get("notes", ""))
    doc.add_heading("Flight Schedule", 1)
    doc.add_paragraph(plan.get("flight_schedule", ""))
    doc.add_heading("Format Requirements", 1)
    doc.add_paragraph("Preferred formats: bulletin (14x48), poster (10.5x36), junior poster (6x12). "
                      "Digital OOH accepted. Provide production specs and posting dates.")
    return _to_bytes(doc)


def generate_rfp_template(deal: dict, plan: dict, vendor_type: str = "print") -> bytes:
    label = vendor_type.upper()
    doc = _new_doc(f"Request for Proposal — {label} — {deal['name']}")

    doc.add_heading("RFP Overview", 1)
    doc.add_paragraph(
        f"Hilco Global is soliciting proposals for {vendor_type} media placements "
        f"in support of the following deal: {deal['name']}."
    )
    _deal_overview(doc, deal)

    if vendor_type not in ("print", "ooh"):
        raise ValueError(f"vendor_type must be 'print' or 'ooh', got: {vendor_type!r}")
    budget_key = vendor_type
    channel_budget = plan["channel_mix"].get(budget_key, {}).get("usd", 0)
    doc.add_heading("Budget", 1)
    _bold_para(doc, "Available Budget", f"${channel_budget:,.0f}")

    doc.add_heading("Required Deliverables", 1)
    doc.add_paragraph("Please include in your proposal:")
    for item in ["Proposed placement(s) with description", "Audience reach and demographics",
                 "Pricing and production costs", "Creative specifications and deadlines",
                 "Sample placements or media kit"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Submission Deadline", 1)
    _bold_para(doc, "Submit by", f"5 business days prior to {deal['start_date']}")
    _bold_para(doc, "Submit to", "[Your Name] — [Your Email]")
    return _to_bytes(doc)


def generate_budget_approval_form(deal: dict, plan: dict) -> bytes:
    doc = _new_doc(f"Media Budget Approval — {deal['name']}")

    doc.add_heading("Deal Information", 1)
    _deal_overview(doc, deal)
    _bold_para(doc, "Total Media Budget Requested", f"${deal['total_budget']:,.0f}")

    doc.add_heading("Channel Allocation", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Channel", "% of Budget", "Amount"
    for channel, data in plan["channel_mix"].items():
        row = table.add_row().cells
        row[0].text = channel.title()
        row[1].text = f"{data['pct']}%"
        row[2].text = f"${data['usd']:,.0f}"

    doc.add_heading("Rationale", 1)
    doc.add_paragraph(plan.get("rationale", ""))

    doc.add_heading("Approval", 1)
    for label in ["Approved by", "Title", "Date", "Signature"]:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run("_" * 40)
    return _to_bytes(doc)


def generate_media_plan_summary(deal: dict, plan: dict) -> bytes:
    doc = _new_doc(f"Media Plan Summary — {deal['name']}")
    _deal_overview(doc, deal)
    _bold_para(doc, "Total Budget", f"${deal['total_budget']:,.0f}")

    doc.add_heading("Channel Mix", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Channel", "%", "Budget", "Notes"
    for channel, data in plan["channel_mix"].items():
        row = table.add_row().cells
        row[0].text = channel.title()
        row[1].text = f"{data['pct']}%"
        row[2].text = f"${data['usd']:,.0f}"
        row[3].text = data.get("notes", "")

    doc.add_heading("Rationale", 1)
    doc.add_paragraph(plan.get("rationale", ""))
    doc.add_heading("Flight Schedule", 1)
    doc.add_paragraph(plan.get("flight_schedule", ""))
    doc.add_heading("Targeting Notes", 1)
    doc.add_paragraph(plan.get("targeting_notes", ""))
    if plan.get("keywords"):
        _bold_para(doc, "Keywords", ", ".join(plan["keywords"]))
    return _to_bytes(doc)


def generate_performance_report(deal: dict, snapshots: list, placements: list) -> bytes:
    """Returns a PDF performance report as bytes."""
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                             rightMargin=inch, leftMargin=inch,
                             topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(f"Performance Report — {deal['name']}", styles["Title"]))
    story.append(Paragraph(
        f"Location: {deal['location']} | Budget: ${deal['total_budget']:,.0f} | "
        f"{deal['start_date']} to {deal['end_date']}",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.2 * inch))

    story.append(Paragraph("Digital Platform Performance", styles["Heading2"]))
    table_data = [["Platform", "Impressions", "Clicks", "Spend", "Conversions", "CPM", "CPC"]]
    total_imp = total_clicks = total_spend = total_conv = 0
    for s in snapshots:
        imp, clk, spd, conv = s["impressions"], s["clicks"], s["spend"], s["conversions"]
        cpm = (spd / imp * 1000) if imp else 0
        cpc = (spd / clk) if clk else 0
        table_data.append([
            s["platform"].title(), f"{imp:,}", f"{clk:,}",
            f"${spd:,.0f}", str(conv), f"${cpm:.2f}", f"${cpc:.2f}",
        ])
        total_imp += imp
        total_clicks += clk
        total_spend += spd
        total_conv += conv
    blend_cpm = (total_spend / total_imp * 1000) if total_imp else 0
    blend_cpc = (total_spend / total_clicks) if total_clicks else 0
    table_data.append([
        "TOTAL", f"{total_imp:,}", f"{total_clicks:,}",
        f"${total_spend:,.0f}", str(total_conv), f"${blend_cpm:.2f}", f"${blend_cpc:.2f}",
    ])

    t = Table(table_data, hAlign="LEFT")
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#d9e2f3")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#f5f5f5")]),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.2 * inch))

    manual = [p for p in placements if not p.get("is_digital", False)]
    if manual:
        story.append(Paragraph("Print & OOH Placements", styles["Heading2"]))
        p_data = [["Vendor", "Channel", "Contracted", "Actual Spend", "Status"]]
        for p in manual:
            p_data.append([
                p["vendor_name"], p["channel"].upper(),
                f"${p['contracted_cost']:,.0f}", f"${p['actual_spend']:,.0f}", p["status"].title(),
            ])
        pt = Table(p_data, hAlign="LEFT")
        pt.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(pt)

    doc.build(story)
    return buf.getvalue()


def generate_creative_brief_doc(deal: dict, plan: dict, creative: dict) -> bytes:
    """Returns a Word doc with all creative ideas for all channels."""
    doc = _new_doc(f"Creative Brief — {deal['name']}")
    _deal_overview(doc, deal)

    # Google Search
    gs = creative.get("google_search", {})
    if gs:
        doc.add_heading("Google Search Ads (RSA)", 1)
        doc.add_paragraph("Headlines (30 chars max each):", style="Intense Quote")
        for h in gs.get("headlines", []):
            doc.add_paragraph(h, style="List Number")
        doc.add_paragraph("Descriptions (90 chars max each):", style="Intense Quote")
        for d in gs.get("descriptions", []):
            doc.add_paragraph(d, style="List Number")
        if gs.get("notes"):
            _bold_para(doc, "Notes", gs["notes"])

    # Meta
    meta = creative.get("meta", {})
    if meta:
        doc.add_heading("Meta (Facebook / Instagram)", 1)
        doc.add_paragraph("Primary Text Options:", style="Intense Quote")
        for t in meta.get("primary_text", []):
            doc.add_paragraph(t, style="List Number")
        doc.add_paragraph("Headlines:", style="Intense Quote")
        for h in meta.get("headlines", []):
            doc.add_paragraph(h, style="List Number")
        _bold_para(doc, "Recommended CTA", meta.get("cta", ""))
        if meta.get("notes"):
            _bold_para(doc, "Notes", meta["notes"])

    # LinkedIn
    li = creative.get("linkedin", {})
    if li:
        doc.add_heading("LinkedIn Sponsored Content", 1)
        doc.add_paragraph("Headline Options:", style="Intense Quote")
        for h in li.get("headline", []):
            doc.add_paragraph(h, style="List Number")
        doc.add_paragraph("Body Copy Options:", style="Intense Quote")
        for b in li.get("body", []):
            doc.add_paragraph(b, style="List Number")
        if li.get("notes"):
            _bold_para(doc, "Notes", li["notes"])

    # Print
    pr = creative.get("print", {})
    if pr:
        doc.add_heading("Print Advertising", 1)
        doc.add_paragraph("Headline Options:", style="Intense Quote")
        for h in pr.get("headlines", []):
            doc.add_paragraph(h, style="List Number")
        _bold_para(doc, "Short Body Copy", pr.get("body_short", ""))
        _bold_para(doc, "Long Body Copy", pr.get("body_long", ""))
        _bold_para(doc, "Call to Action", pr.get("cta", ""))
        if pr.get("notes"):
            _bold_para(doc, "Notes", pr["notes"])

    # OOH
    ooh = creative.get("ooh", {})
    if ooh:
        doc.add_heading("Out-of-Home (Billboards / Signage)", 1)
        doc.add_paragraph("Headline Options (5-7 words max):", style="Intense Quote")
        for h in ooh.get("headlines", []):
            doc.add_paragraph(h, style="List Number")
        doc.add_paragraph("Subheadlines:", style="Intense Quote")
        for s in ooh.get("subheads", []):
            doc.add_paragraph(s, style="List Number")
        _bold_para(doc, "Call to Action", ooh.get("cta", ""))
        if ooh.get("notes"):
            _bold_para(doc, "Notes", ooh["notes"])

    return _to_bytes(doc)
